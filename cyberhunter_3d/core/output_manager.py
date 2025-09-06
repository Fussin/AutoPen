import json
import logging
from pathlib import Path
from datetime import datetime
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class OutputManager:
    def __init__(self, base_dir):
        self.base_dir = Path(base_dir)
        self.recon_dir = self.base_dir / "recon"
        self.network_dir = self.base_dir / "network"
        self.discovery_dir = self.base_dir / "discovery"
        self.reports_dir = self.base_dir / "reports"

        self.vulnerabilities = []
        self.assets = []

        self.base_dir.mkdir(exist_ok=True)
        self.recon_dir.mkdir(exist_ok=True)
        self.network_dir.mkdir(exist_ok=True)
        self.discovery_dir.mkdir(exist_ok=True)
        self.reports_dir.mkdir(exist_ok=True)

    @classmethod
    def create_for_timestamp(cls, base_path):
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        scan_dir = Path(base_path) / timestamp
        return cls(scan_dir)

    def write_recon_file(self, filename, content):
        (self.recon_dir / filename).write_text(content)

    def write_network_json(self, filename, data):
        with (self.network_dir / filename).open('w') as f:
            json.dump(data, f, indent=2)

    def write_discovery_file(self, filename, content):
        (self.discovery_dir / filename).write_text(content)

    def add_vulnerability(self, vuln_data, severity=""):
        vuln_data['severity'] = severity
        self.vulnerabilities.append(vuln_data)

    def add_asset(self, asset_data):
        self.assets.append(asset_data)

    def produce_metadata(self):
        metadata = {
            "timestamp": datetime.now().isoformat(),
            "total_assets": len(self.assets),
            "total_vulnerabilities": len(self.vulnerabilities),
        }
        with (self.base_dir / "metadata.json").open('w') as f:
            json.dump(metadata, f, indent=2)

    def finalize(self, generate_pdf=False, generate_docx=False):
        summary = {
            "total_assets": len(self.assets),
            "vulnerabilities": len(self.vulnerabilities),
            "reports": []
        }

        if generate_pdf:
            pdf_path = self._generate_pdf_report()
            summary["reports"].append({"type": "pdf", "path": str(pdf_path)})

        if generate_docx:
            docx_path = self._generate_docx_report()
            summary["reports"].append({"type": "docx", "path": str(docx_path)})

        return summary

    def _generate_pdf_report(self):
        logger.info("Generating PDF report...")
        template_env = Environment(loader=FileSystemLoader("cyberhunter_3d/reporting/templates/"))
        template = template_env.get_template("report.html")

        context = {
            "scan_id": self.base_dir.name,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "summary": {
                "total_assets": len(self.assets),
                "vulnerabilities": len(self.vulnerabilities),
            },
            "vulnerabilities": self.vulnerabilities,
            "assets": self.assets,
        }

        html_out = template.render(context)
        pdf_path = self.reports_dir / "scan_report.pdf"
        HTML(string=html_out).write_pdf(pdf_path)
        logger.info(f"PDF report generated at: {pdf_path}")
        return pdf_path

    def _generate_docx_report(self):
        logger.info("Generating DOCX report...")
        doc = Document()
        doc.add_heading('Scan Report', 0)
        doc.add_paragraph(f"Scan ID: {self.base_dir.name}")
        doc.add_paragraph(f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f"Total assets found: {len(self.assets)}")
        doc.add_paragraph(f"Vulnerabilities found: {len(self.vulnerabilities)}")

        doc.add_heading('Vulnerabilities', level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Title'
        hdr_cells[2].text = 'Severity'
        hdr_cells[3].text = 'Description'

        for vuln in self.vulnerabilities:
            row_cells = table.add_row().cells
            row_cells[0].text = str(vuln.get('id', 'N/A'))
            row_cells[1].text = vuln.get('title', '')
            row_cells[2].text = vuln.get('severity', '')
            row_cells[3].text = vuln.get('description', '')

        doc.add_heading('Discovered Assets', level=1)
        asset_table = doc.add_table(rows=1, cols=3)
        asset_hdr_cells = asset_table.rows[0].cells
        asset_hdr_cells[0].text = 'Type'
        asset_hdr_cells[1].text = 'Value'
        asset_hdr_cells[2].text = 'Details'

        for asset in self.assets:
            row_cells = asset_table.add_row().cells
            row_cells[0].text = asset.get('type', '')
            row_cells[1].text = asset.get('value', '')
            row_cells[2].text = json.dumps(asset.get('details', {}))

        docx_path = self.reports_dir / "scan_report.docx"
        doc.save(docx_path)
        logger.info(f"DOCX report generated at: {docx_path}")
        return docx_path
